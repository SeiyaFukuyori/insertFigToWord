import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

doc = docx.Document('./input/画像埋め込みテスト.docx')
tbl = doc.tables[0]
values = []
row_index = len(tbl.rows) - 1
col_index = len(tbl.columns) - 1
p = tbl.rows[row_index].cells[col_index].paragraphs[0]
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture('./fig/test.png', width=Cm(8.0), height=Cm(6.0))
r.add_break()
doc.save('./result/inserted_result.docx')